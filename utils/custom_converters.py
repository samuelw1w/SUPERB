import docx
from typing import List, Union, Dict, Any, Optional
from haystack import Document
from pathlib import Path

class DocxToTextConverter:
    """
    A component to convert docx file to Document
    """
    def run(
        self,
        sources: List[Union[str, Path]],
        meta: Optional[Union[Dict[str, Any], List[Dict[str, Any]]]] = None,
    ):
        if meta is None:
            meta = {}
        documents = []
        for file_path in sources:
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            doc = Document(content=text, meta=meta)
            documents.append(doc)
        return {"documents": documents}
